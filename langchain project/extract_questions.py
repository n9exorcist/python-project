import os
import time
from google import genai
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

# ─────────────────────────────────────────────
# CONFIGURATION — edit these two lines only
# ─────────────────────────────────────────────
FOLDER_PATH = r"C:\Users\narayanan.selvaraj\Downloads\files"   # ← your image folder
OUTPUT_FILE = r"C:\Users\narayanan.selvaraj\Downloads\files\Compiled_Questions.docx"
# ─────────────────────────────────────────────

load_dotenv()
api_key = os.getenv("GEMINI_API_KEY")

if not api_key:
    raise ValueError("GEMINI_API_KEY not found. Please check your .env file.")

client = genai.Client(api_key=api_key)

# ── Word document setup ──────────────────────
doc = Document()

# Title
title = doc.add_heading('Compiled MCQ Questions', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
sub = doc.add_paragraph(f'Auto-extracted from scanned images')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
sub.runs[0].font.size = Pt(10)

doc.add_paragraph()  # spacer

# ── Extraction settings ──────────────────────
seen_questions = set()
question_number = 0
file_count = 0
skipped = 0
MAX_RETRIES = 5

PROMPT = """
You are extracting a multiple-choice question from an exam paper image.

Extract ONLY the question and its options. Format EXACTLY as:
Question: [The full question text]
A) [Option A text]
B) [Option B text]
C) [Option C text]
D) [Option D text]

Rules:
- Do NOT include any answer key or explanation
- Do NOT add extra commentary
- Preserve the exact wording from the image
- If the image has no MCQ, reply with: NO_QUESTION
"""

print(f"Scanning folder: {FOLDER_PATH}")
image_files = sorted([f for f in os.listdir(FOLDER_PATH)
                      if f.lower().endswith(('.jpg', '.jpeg', '.png'))])
total_files = len(image_files)
print(f"Found {total_files} image files.\n")

for filename in image_files:
    file_path = os.path.join(FOLDER_PATH, filename)
    file_count += 1
    print(f"[{file_count}/{total_files}] Processing: {filename}")

    sample_file = None
    success = False

    for attempt in range(MAX_RETRIES):
        try:
            sample_file = client.files.upload(file=file_path)

            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=[sample_file, PROMPT]
            )

            extracted_text = response.text.strip()

            # Skip if Gemini says no question found
            if "NO_QUESTION" in extracted_text.upper():
                print(f"  → No MCQ found in image, skipping.")
                success = True
                break

            # Deduplicate on first 50 chars of question line
            question_line = ""
            for line in extracted_text.splitlines():
                if line.lower().startswith("question:"):
                    question_line = line[9:].strip().lower().replace(" ", "")[:50]
                    break

            dedup_key = question_line or extracted_text[:50].lower().replace(" ", "")

            if dedup_key in seen_questions:
                print(f"  → Duplicate detected, skipping.")
                skipped += 1
                success = True
                break

            seen_questions.add(dedup_key)
            question_number += 1

            # ── Write to Word doc ────────────────
            # Question number header
            num_para = doc.add_paragraph()
            num_run = num_para.add_run(f"Question {question_number}")
            num_run.bold = True
            num_run.font.size = Pt(11)
            num_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)  # dark blue

            # Content lines
            for line in extracted_text.splitlines():
                line = line.strip()
                if not line:
                    continue
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                # Bold the "Question:" label and option letters
                if line.lower().startswith("question:") or (
                    len(line) >= 2 and line[0].upper() in "ABCD" and line[1] == ")"
                ):
                    run.bold = (line[0].upper() in "ABCD" and line[1] == ")")
                p.paragraph_format.space_after = Pt(2)

            # Separator line
            sep = doc.add_paragraph("─" * 60)
            sep.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            sep.runs[0].font.size = Pt(9)
            doc.add_paragraph()  # breathing room

            print(f"  → Added as Question {question_number} ✓")
            success = True
            break

        except Exception as e:
            error_msg = str(e)
            if '503' in error_msg or '429' in error_msg or 'UNAVAILABLE' in error_msg:
                wait_time = 5 * (attempt + 1)
                print(f"  → Server busy. Retrying in {wait_time}s (attempt {attempt + 1}/{MAX_RETRIES})")
                time.sleep(wait_time)
            else:
                print(f"  → Unexpected error: {e}")
                break

        finally:
            # Always clean up uploaded file
            if sample_file:
                try:
                    client.files.delete(name=sample_file.name)
                except Exception:
                    pass

    if not success:
        print(f"  → FAILED after {MAX_RETRIES} attempts: {filename}")

# ── Final summary page ───────────────────────
doc.add_page_break()
summary = doc.add_heading('Summary', level=1)
doc.add_paragraph(f'Total images scanned:   {total_files}')
doc.add_paragraph(f'Unique questions found: {question_number}')
doc.add_paragraph(f'Duplicates skipped:     {skipped}')
doc.add_paragraph(f'Failed to process:      {total_files - file_count + (MAX_RETRIES if not success else 0)}')

doc.save(OUTPUT_FILE)
print(f"\n{'='*50}")
print(f"Done! {question_number} unique questions saved to:")
print(f"  {OUTPUT_FILE}")
print(f"Duplicates skipped: {skipped}")
print(f"{'='*50}")